import os
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from langchain_community.document_loaders import TextLoader
from langchain_community.document_loaders.csv_loader import CSVLoader


def list_txt_files(data_dir="./data"):
    paths = Path(data_dir).glob('**/*.txt')
    for path in paths:
        yield str(path)


def load_txt_files(data_dir="./data"):
    docs = []
    paths = list_txt_files(data_dir)
    for path in paths:
        print(f"Loading {path}")
        loader = TextLoader(path)
        docs.extend(loader.load())
    return docs


def load_csv_files(data_dir="./data"):
    docs = []
    paths = Path(data_dir).glob('**/*.csv')
    for path in paths:
        loader = CSVLoader(file_path=str(path))
        docs.extend(loader.load())
    return docs


def load_pdf_files(data_dir="./data"):
    docs = []
    paths = Path(data_dir).glob('**/*.pdf')
    for path in paths:
        print(f"Loading {path}")
        pdf_reader = PdfReader(path)
        for num, page in enumerate(pdf_reader.pages):
            page_content = page.extract_text()
            doc = Document(page_content=page_content, metadata={'title': str(path), 'page': num + 1})
            docs.append(doc)
    return docs


def load_docx_files(data_dir="./data"):
    docs = []
    paths = Path(data_dir).glob('**/*.docx')
    for path in paths:
        print(f"Loading {path}")
        docx = DocxDocument(path)
        for paragraph in docx.paragraphs:
            if paragraph.text.strip():
                doc = Document(page_content=paragraph.text, metadata={'title': str(path)})
                docs.append(doc)
    return docs


# Use with result of file_to_summarize = st.file_uploader("Choose a file") or a string.
# or a file-like object.
def get_document_text(uploaded_file, title=None):
    docs = []
    fname = uploaded_file.name
    if not title:
        title = os.path.basename(fname)
    if fname.lower().endswith('pdf'):
        pdf_reader = PdfReader(uploaded_file)
        for num, page in enumerate(pdf_reader.pages):
            page = page.extract_text()
            doc = Document(page_content=page, metadata={'title': title, 'page': (num + 1)})
            docs.append(doc)
    elif fname.lower().endswith('docx'):
        docx = DocxDocument(uploaded_file)
        for paragraph in docx.paragraphs:
            if paragraph.text.strip():
                doc = Document(page_content=paragraph.text, metadata={'title': title})
                docs.append(doc)
    else:
        # assume text
        doc_text = uploaded_file.read().decode()
        docs.append(doc_text)

    return docs
from typing import List


def get_document_text_2(folder_path: str) -> List[Document]:
    """
    Processes all PDF, DOCX, and TXT files in the given folder path.

    Args:
        folder_path (str): Path to the folder containing the documents.

    Returns:
        List[Document]: A list of Document objects with text and metadata.
    """
    docs = []

    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            title = os.path.basename(file_path)

            if file.lower().endswith('pdf'):
                try:
                    pdf_reader = PdfReader(file_path)
                    for num, page in enumerate(pdf_reader.pages):
                        page_content = page.extract_text()
                        if page_content.strip():
                            doc = Document(page_content=page_content, metadata={'title': title, 'page': (num + 1)})
                            docs.append(doc)
                except Exception as e:
                    print(f"Error processing PDF {file_path}: {e}")

            elif file.lower().endswith('docx'):
                try:
                    docx = DocxDocument(file_path)
                    for paragraph in docx.paragraphs:
                        if paragraph.text.strip():
                            doc = Document(page_content=paragraph.text, metadata={'title': title})
                            docs.append(doc)
                except Exception as e:
                    print(f"Error processing DOCX {file_path}: {e}")

            elif file.lower().endswith('txt'):
                try:
                    with open(file_path, 'r', encoding='utf-8') as txt_file:
                        doc_text = txt_file.read().strip()
                        if doc_text:
                            doc = Document(page_content=doc_text, metadata={'title': title})
                            docs.append(doc)
                except Exception as e:
                    print(f"Error processing TXT {file_path}: {e}")

            else:
                print(f"Unsupported file type: {file_path}")

    return docs



if __name__ == "__main__":
    # Example usage for single files
    # example_pdf_path = "examples/healthy_meal_10_tips.pdf"
    # docs = get_document_text(open(example_pdf_path, "rb"))
    # for doc in docs:
    #     print(doc)
    # docs = get_document_text(open("examples/us_army_recipes.txt", "rb"))
    # for doc in docs:
    #     print(doc)

    # Example usage for directory processing
    print("Processing all TXT files:")
    txt_docs = load_txt_files("examples")
    for doc in txt_docs:
        print(doc)

    print("Processing all CSV files:")
    csv_docs = load_csv_files("examples")
    for doc in csv_docs:
        print(doc)

    print("Processing all PDF files:")
    pdf_docs = load_pdf_files("examples")
    for doc in pdf_docs:
        print(doc)

    print("Processing all DOCX files:")
    docx_docs = load_docx_files("examples")
    for doc in docx_docs:
        print(doc)
